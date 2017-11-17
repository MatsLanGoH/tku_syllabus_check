# encoding: utf-8
import os
import sys

# Excel workbook handling
from xlrd import open_workbook

# Word document handling
from docx import Document
from docx.shared import Inches, RGBColor

# Data positions in excel workbooks
from positions import kamoku_info, kamoku_maps

# Check instructions
from instructions import checks
from check import Check

# Constants
RESULTS_FOLDER = u"results"


def check_sheet(sheet):
    """
    Checks a single worksheet.
    """
    # Step 1: Setup environment
    k_items = kamoku_maps.get(sheet.nrows)

    warnings = ""
    # Step 2: Execute checks in checks.
    # NOTE: Make sure we don't have to recreate the check_series each time.
    for check in checks:
        ck = Check(check[0], check[1], k_items, sheet)
        if not ck.passes():
            warnings += '\n□ ' + ck.msg
            warnings += '\n' + ck.errors

    # print(warnings)
    return warnings


def execute_checks(workbook):
    """
    Executes checks for a given workbook.
    Saves result to a docx file containing the name of the source workbook
    as its file name.
    """
    # Create output directory if necessary
    if not os.path.exists(os.path.join(path, RESULTS_FOLDER)):
        os.makedirs(os.path.join(path, RESULTS_FOLDER))

    # Open workbook and get its name to create output file name
    xl_workbook = open_workbook(workbook)
    result_name = workbook.split('/')[-1]
    result_name = result_name.split('.')[0] + '-tenken.docx'
    result_location = os.path.join(path, RESULTS_FOLDER, result_name)

    # Setup output file
    document = Document()
    document.add_heading(u'シラバス点検 自動点検結果', 0)

    # Get the sheets from the workbook
    sheets = xl_workbook.sheets()

    # Execute checks for each sheet
    for sheet in sheets:
        # Build strings for result output
        teacher_str = u'\n\t'.join(sheet.cell_value(*kamoku_info['teacher'])
                                   .split('\n'))
        class_str = sheet.cell_value(*kamoku_info['name'])

        # Get check warnings
        warnings = check_sheet(sheet)

        # Add warnings to document
        document.add_heading('{}: {}'.format(sheet.name, teacher_str), level=2)
        p = document.add_paragraph()
        p.add_run('授業科目名: {}'.format(class_str)).bold = True
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run('{}'.format(warnings))

        # Format colors according to warnings status
        if warnings:
            p.add_run('\n => 上記の項目を確認し、必要な場合は修正を行ってください。')\
                    .font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        else:
            p.add_run('\n => 点検OK。')
        p.add_run('\n' + '_' * 78 + '\n')  # Break line

    # Finally, save result file
    document.save(result_location)


def main(path):
    """
    Finds all Workbooks in directory and checks the syllabus.
    """
    # Read workbooks from directory
    workbooks = [os.path.join(path, f) for f in os.listdir(path) 
                 if os.path.isfile(os.path.join(path, f)) and
                 f.lower().endswith('.xls')]
    print(workbooks)

    # TODO: Do we want to escape like this?
    if not workbooks:
        print('No workbooks in directory.')
        return

    # Check syllabus for all workbooks.
    for wb in workbooks:
        execute_checks(wb)


if __name__ == "__main__":
    if len(sys.argv) >= 2:
        path = os.path.join(os.getcwd(), os.path.normpath(sys.argv[1]))
    else:
        path = os.getcwd()

    main(path)
